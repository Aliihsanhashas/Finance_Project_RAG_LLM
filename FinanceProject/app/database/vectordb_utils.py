
import os 

from bs4 import BeautifulSoup

from docx import Document

from docx import Document




def read_doc(doc_path):
    with open(doc_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    # Parse the HTML content
    soup = BeautifulSoup(content, "html.parser")

    result = []

    for element in soup.body.find_all(recursive=False):  # Process direct children in order
        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]:
            result.append(element.get_text(strip=True))
        elif element.name == "table":
            table_text = []
            for row in element.find_all("tr"):
                row_text = [cell.get_text(strip=True) for cell in row.find_all(["td", "th"])]
                table_text.append("\t".join(row_text))  # Tab-separated row
            result.append("\n".join(table_text))  # Join rows to form a readable table
    
    return "\n\n".join(result)  # Join all content with spacing





# **1. `.docx` Dosyasını Oku (Tek Parça Olarak)**
def read_word_file(file_path):
    """Word dosyasını okuyup içindeki metni ve tabloları tek parça olarak döndürür."""

    if os.path.splitext(file_path)[1] == ".doc": 
        return read_doc(file_path)

    else:
            
        doc = Document(file_path)
        text = []
    
        # Paragrafları ekleyelim
        for para in doc.paragraphs:
            clean_text = para.text.strip()
            if clean_text:  # Boş paragrafları ekleme
                text.append(clean_text)
    
        # Tabloları ekleyelim
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):  # Eğer satır tamamen boş değilse ekleyelim
                    text.append(" | ".join(row_text))  # Satırları '|' ile ayırarak ekleyelim
    
        return "\n".join(text)  # Satırları birleştirerek tek parça olarak döndür




def write_to_vectordb(file_path, collection,  metadata, chunk_size=1000):
    # **2. ChromaDB İçin Embedding Modelini Tanımla**
    
    # **3. `.docx` Dosyasını Yükle (Parçalama Yapmadan)**
    
    #import pdb;pdb.set_trace()
    document_text = read_word_file(file_path)
    
    # **4. Veriyi parça parça Olarak ChromaDB'ye Kaydet**
    for chunk_part in range(0, len(document_text), chunk_size):
        collection.add(documents = [document_text[chunk_part:chunk_part + chunk_size] ], metadatas = [metadata ], ids = [str(random.randint(0, 10000000))])
    print("Veriler tek parça olarak ChromaDB'ye kaydedildi!")
    
